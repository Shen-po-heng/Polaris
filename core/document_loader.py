"""Document loader — unified loader for PDF, DOCX, TXT, and MD files.

Each loaded document is a LangChain ``Document`` object with ``page_content``
and ``metadata`` fields.  Metadata always includes:
  - ``source``: filename (no directory path, for privacy)
  - ``page``:   page number for PDFs, 0 for single-chunk formats

Usage:
    from core.document_loader import DocumentLoader
    docs = DocumentLoader.load("/path/to/paper.pdf")
"""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader

from core.exceptions import DocumentLoadError
from utils.logger import get_logger
from utils.security import validate_file

logger = get_logger(__name__)

# Formats that can be read as plain text without a special library.
_PLAIN_TEXT_SUFFIXES: frozenset[str] = frozenset({".txt", ".md"})


class DocumentLoader:
    """Unified document loader — PDF / DOCX / TXT / MD."""

    @staticmethod
    def load(file_path: str | Path) -> list[Document]:
        """Load a single document file into a list of LangChain Documents.

        For PDFs each page becomes its own Document.
        For DOCX, TXT, and MD the entire file is one Document.

        Args:
            file_path: Absolute path to the document.

        Returns:
            Non-empty list of Documents.

        Raises:
            DocumentLoadError: If the file cannot be read or parsed.
        """
        validated = validate_file(file_path)
        suffix = validated.suffix.lower()

        try:
            if suffix == ".pdf":
                docs = DocumentLoader._load_pdf(validated)
            elif suffix == ".docx":
                docs = DocumentLoader._load_docx(validated)
            elif suffix in _PLAIN_TEXT_SUFFIXES:
                docs = DocumentLoader._load_text(validated)
            else:
                raise DocumentLoadError(f"Unsupported file type: {suffix!r}")
        except DocumentLoadError:
            raise
        except Exception as exc:
            logger.error("Failed to load '%s': %s", validated.name, exc)
            raise DocumentLoadError(f"Cannot load '{validated.name}': {exc}") from exc

        if not docs:
            raise DocumentLoadError(f"'{validated.name}' produced no content.")

        logger.info("Loaded '%s' — %d document(s)", validated.name, len(docs))
        return docs

    # ── Format-specific loaders ─────────────────────────────────────────────

    @staticmethod
    def _load_pdf(path: Path) -> list[Document]:
        loader = PyPDFLoader(str(path))
        docs = loader.load()
        for doc in docs:
            doc.metadata["source"] = path.name
            doc.metadata.setdefault("page", 0)
        return docs

    @staticmethod
    def _load_docx(path: Path) -> list[Document]:
        # python-docx is an optional dependency; import lazily so missing
        # package raises a clear error at call time rather than import time.
        try:
            import docx  # type: ignore[import-untyped]
        except ImportError as exc:
            raise DocumentLoadError(
                "python-docx is required for .docx files: pip install python-docx"
            ) from exc

        word_doc = docx.Document(str(path))
        text = "\n".join(para.text for para in word_doc.paragraphs if para.text.strip())
        return [Document(page_content=text, metadata={"source": path.name, "page": 0})]

    @staticmethod
    def _load_text(path: Path) -> list[Document]:
        text = path.read_text(encoding="utf-8", errors="replace")
        return [Document(page_content=text, metadata={"source": path.name, "page": 0})]
